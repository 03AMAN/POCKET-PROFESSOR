from flask import Blueprint, render_template, request, send_file, redirect, url_for
import pdfkit
import os
import sqlite3
from docx import Document
from io import BytesIO

srb_bp = Blueprint('smart-resume', __name__,template_folder='../templates')

# Path to wkhtmltopdf executable
pdfkit_config = pdfkit.configuration(wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')

# SQLite database name
DB_NAME = 'resumes.db'

def init_db():
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute(''' 
            CREATE TABLE IF NOT EXISTS resumes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT, email TEXT, phone TEXT,
                skills TEXT, experience TEXT, education TEXT,
                certifications TEXT, projects TEXT , achievements TEXT
            )
        ''')

@srb_bp.route('/srb')
def index():
    return render_template('form.html')

@srb_bp.route('/srb/preview', methods=['POST'])
def preview():
    data = request.form.to_dict()
    template = request.form.get('template')
    return render_template(template, data=data)

@srb_bp.route('/srb/save', methods=['POST'])
def save():
    data = request.form.to_dict()
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute('''
            INSERT INTO resumes (name, email, phone, skills, experience, education, certifications, projects, achievements)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (data['name'], data['email'], data['phone'],
              data['skills'], data['experience'], data['education'], data['certifications'], data['projects'], data['achievements']))
    return redirect(url_for('smart-resume.list_resumes'))  # Updated to use the correct endpoint

@srb_bp.route('/srb/list')
def list_resumes():
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.execute("SELECT * FROM resumes")
        resumes = cursor.fetchall()
    return render_template('list.html', resumes=resumes)

@srb_bp.route('/srb/download/<int:resume_id>/<format>')
def download(resume_id, format):
    with sqlite3.connect(DB_NAME) as conn:
        cursor = conn.execute("SELECT * FROM resumes WHERE id=?", (resume_id,))
        row = cursor.fetchone()
        if not row:
            return "Resume not found", 404

    data = {
        'name': row[1], 'email': row[2], 'phone': row[3],
        'skills': row[4], 'experience': row[5], 'education': row[6],
        'certifications': row[7], 'projects': row[8], 'achievements': row[9]
    }

    if format == 'pdf':
        rendered = render_template('template1.html', data=data)
        options = {'enable-local-file-access': None}
        pdf = pdfkit.from_string(rendered, False, configuration=pdfkit_config, options=options)
        return send_file(BytesIO(pdf), as_attachment=True, download_name="resume.pdf")

    elif format == 'word':
        doc = Document()
        doc.add_heading(data['name'], 0)
        doc.add_paragraph(f"Email: {data['email']}")
        doc.add_paragraph(f"Phone: {data['phone']}")
        doc.add_paragraph(f"Skills: {data['skills']}")
        doc.add_paragraph(f"Experience: {data['experience']}")
        doc.add_paragraph(f"Education: {data['education']}")
        doc.add_paragraph(f"Certifications: {data['certifications']}")
        doc.add_paragraph(f"Projects: {data['projects']}")
        doc.add_paragraph(f"Achievements: {data['achievements']}")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="resume.docx")

    return "Invalid format", 400

@srb_bp.route('/srb/delete/<int:resume_id>', methods=['POST'])
def delete_resume(resume_id):
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("DELETE FROM resumes WHERE id=?", (resume_id,))
    return redirect(url_for('smart-resume.list_resumes'))  # Updated to use the correct endpoint

# Optional: call this from app.py
def init_srb_db():
    init_db()
